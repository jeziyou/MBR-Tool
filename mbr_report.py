"""
MBR 工艺计算书 - PDF / Word 生成模块
"""
import io
from datetime import datetime
from mbr_calc import ProcessInput, ProcessResult

# ============================================================================
# 辅助格式化函数
# ============================================================================
def _f(v, decimals=1):
    if v is None:
        return "—"
    return f"{v:.{decimals}f}"


def _fi(v):
    if v is None:
        return "—"
    return f"{round(v):,}"


# ============================================================================
# PDF 生成（reportlab）
# ============================================================================
def generate_pdf_report(inp: ProcessInput, result: ProcessResult,
                        project_name: str = "MBR 膜系统工艺计算书",
                        designer: str = "", design_date: str = "") -> bytes:
    """生成 PDF 计算书，返回 bytes"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    )
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # 注册中文字体，按优先级尝试
    font_name = "Helvetica"
    font_bold = "Helvetica-Bold"
    font_cjk_candidates = [
        ("STSong", "/System/Library/Fonts/STSong.ttc"),
        ("STSong", "/usr/share/fonts/truetype/arphic/uming.ttc"),
        ("NotoSansCJK", "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        ("NotoSerifCJK", "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc"),
        ("WenQuanYi", "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
        ("WenQuanYiMicro", "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"),
        ("DejaVu", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        ("SimSun", "C:/Windows/Fonts/simsun.ttc"),
        ("MicrosoftYaHei", "C:/Windows/Fonts/msyh.ttc"),
    ]
    for fname, fpath in font_cjk_candidates:
        try:
            pdfmetrics.registerFont(TTFont(fname, fpath))
            font_name = fname
            font_bold = fname  # 用同一个字体当作粗体，简化处理
            break
        except Exception:
            continue

    # 构建文档
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
        title=project_name,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title_CN", parent=styles["Title"],
        fontName=font_bold, fontSize=20, leading=28,
        alignment=1, textColor=colors.HexColor("#1f2937"),
    )
    h1_style = ParagraphStyle(
        "H1", parent=styles["Heading1"],
        fontName=font_bold, fontSize=15, leading=22,
        textColor=colors.HexColor("#1e40af"), spaceAfter=10,
    )
    h2_style = ParagraphStyle(
        "H2", parent=styles["Heading2"],
        fontName=font_bold, fontSize=12, leading=18,
        textColor=colors.HexColor("#1f2937"), spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "BodyCN", parent=styles["BodyText"],
        fontName=font_name, fontSize=10, leading=15,
    )
    cell_style = ParagraphStyle(
        "CellCN", parent=styles["BodyText"],
        fontName=font_name, fontSize=9, leading=12, alignment=1,
    )
    header_style = ParagraphStyle(
        "CellHeader", parent=styles["BodyText"],
        fontName=font_bold, fontSize=9, leading=12, alignment=1,
        textColor=colors.white,
    )
    meta_style = ParagraphStyle(
        "Meta", parent=styles["BodyText"],
        fontName=font_name, fontSize=10, leading=15, alignment=1,
        textColor=colors.HexColor("#6b7280"),
    )

    story = []

    # ---- 封面信息 ----
    story.append(Spacer(1, 30))
    story.append(Paragraph(project_name, title_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("MBR 膜生物反应器 - 工艺设计计算书", h2_style))
    story.append(Spacer(1, 20))

    meta_lines = []
    if designer:
        meta_lines.append(f"设计人：{designer}")
    if design_date:
        meta_lines.append(f"设计日期：{design_date}")
    meta_lines.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    for line in meta_lines:
        story.append(Paragraph(line, meta_style))

    story.append(Spacer(1, 30))

    # ---- 章节 1: 设计输入 ----
    story.append(Paragraph("一、设计输入参数", h1_style))

    # 基本参数表
    story.append(Paragraph("1.1 设计流量", h2_style))
    data1 = [
        [Paragraph("参数", header_style), Paragraph("取值", header_style), Paragraph("单位", header_style)],
        [Paragraph("设计流量 Q", cell_style), Paragraph(_fi(inp.Q), cell_style), Paragraph("m³/d", cell_style)],
        [Paragraph("总变化系数 Kz", cell_style), Paragraph(_f(inp.Kz, 2), cell_style), Paragraph("—", cell_style)],
    ]
    _add_table(story, data1)

    story.append(Spacer(1, 10))

    # 进水水质
    story.append(Paragraph("1.2 进水水质", h2_style))
    data2 = [
        [Paragraph("项目", header_style), Paragraph("数值", header_style), Paragraph("单位", header_style)],
        [Paragraph("COD", cell_style), Paragraph(_f(inp.cod_in, 0), cell_style), Paragraph("mg/L", cell_style)],
        [Paragraph("BOD₅", cell_style), Paragraph(_f(inp.bod_in, 0), cell_style), Paragraph("mg/L", cell_style)],
        [Paragraph("NH₃-N", cell_style), Paragraph(_f(inp.nh3n_in, 0), cell_style), Paragraph("mg/L", cell_style)],
        [Paragraph("SS", cell_style), Paragraph(_f(inp.ss_in, 0), cell_style), Paragraph("mg/L", cell_style)],
        [Paragraph("TN", cell_style), Paragraph(_f(inp.tn_in, 0), cell_style), Paragraph("mg/L", cell_style)],
        [Paragraph("TP", cell_style), Paragraph(_f(inp.tp_in, 0), cell_style), Paragraph("mg/L", cell_style)],
        [Paragraph("pH", cell_style), Paragraph(_f(inp.ph_value, 1), cell_style), Paragraph("—", cell_style)],
        [Paragraph("水温", cell_style), Paragraph(_f(inp.T, 1), cell_style), Paragraph("℃", cell_style)],
        [Paragraph("MLSS", cell_style), Paragraph(_f(inp.MLSS, 0), cell_style), Paragraph("mg/L", cell_style)],
    ]
    _add_table(story, data2)

    story.append(Spacer(1, 10))

    # 膜组件配置
    story.append(Paragraph("1.3 膜组件配置", h2_style))
    sheets_opts_map = {
        (40,): "5-53片", (25,): "5-53片", (40, 25): "5-53片",
        (15,): "5-40片", (6,): "5-20片",
    }
    from mbr_calc import ALL_MODELS
    model = ALL_MODELS[inp.model_index]
    data3 = [
        [Paragraph("参数", header_style), Paragraph("取值", header_style), Paragraph("单位", header_style)],
        [Paragraph("膜元件型号", cell_style), Paragraph(model.name, cell_style), Paragraph("—", cell_style)],
        [Paragraph("单膜片面积", cell_style), Paragraph(_f(model.sheet_area, 0), cell_style), Paragraph("m²", cell_style)],
        [Paragraph("膜孔径", cell_style), Paragraph(_f(model.pore_size, 2), cell_style), Paragraph("μm", cell_style)],
        [Paragraph("每架膜片数", cell_style), Paragraph(_fi(inp.sheets_per_rack), cell_style), Paragraph("片/架", cell_style)],
        [Paragraph("池数", cell_style), Paragraph(_fi(inp.pools), cell_style), Paragraph("座", cell_style)],
        [Paragraph("每池架数", cell_style), Paragraph(_fi(inp.racks_per_pool), cell_style), Paragraph("架/池", cell_style)],
    ]
    _add_table(story, data3)

    story.append(Spacer(1, 10))

    # 运行参数
    story.append(Paragraph("1.4 运行与清洗参数", h2_style))
    data4 = [
        [Paragraph("参数", header_style), Paragraph("取值", header_style), Paragraph("单位", header_style)],
        [Paragraph("基准通量 (25℃)", cell_style), Paragraph(_f(inp.J25, 1), cell_style), Paragraph("LMH", cell_style)],
        [Paragraph("污染系数", cell_style), Paragraph(_f(inp.fouling_factor, 2), cell_style), Paragraph("—", cell_style)],
        [Paragraph("比曝气密度 SAD", cell_style), Paragraph(_f(inp.SAD, 0), cell_style), Paragraph("m³/(m²·h)", cell_style)],
        [Paragraph("抽吸 开/停", cell_style), Paragraph(f"{_f(inp.suction_on, 1)} / {_f(inp.suction_off, 1)}", cell_style), Paragraph("min", cell_style)],
        [Paragraph("池内液位", cell_style), Paragraph(_f(inp.pool_level, 1), cell_style), Paragraph("m", cell_style)],
        [Paragraph("管路损失", cell_style), Paragraph(_f(inp.pipe_loss, 1), cell_style), Paragraph("m", cell_style)],
        [Paragraph("产水泵扬程", cell_style), Paragraph(_f(inp.permeate_pump_head, 1), cell_style), Paragraph("m", cell_style)],
        [Paragraph("产水泵效率", cell_style), Paragraph(_f(inp.permeate_pump_eff * 100, 0), cell_style), Paragraph("%", cell_style)],
        [Paragraph("回流比", cell_style), Paragraph(_f(inp.return_ratio, 1), cell_style), Paragraph("—", cell_style)],
        [Paragraph("回流泵扬程", cell_style), Paragraph(_f(inp.return_pump_head, 1), cell_style), Paragraph("m", cell_style)],
        [Paragraph("回流泵效率", cell_style), Paragraph(_f(inp.return_pump_eff * 100, 0), cell_style), Paragraph("%", cell_style)],
        [Paragraph("风机效率", cell_style), Paragraph(_f(inp.fan_efficiency * 100, 0), cell_style), Paragraph("%", cell_style)],
    ]
    _add_table(story, data4)

    story.append(PageBreak())

    # ---- 章节 2: 计算结果 ----
    story.append(Paragraph("二、工艺计算结果", h1_style))

    # 膜面积与通量
    story.append(Paragraph("2.1 膜面积与通量", h2_style))
    flux_status = "✓ 在范围内" if result.j_peak <= result.j_design else "⚠ 超过设计通量"
    data5 = [
        [Paragraph("参数", header_style), Paragraph("结果", header_style), Paragraph("单位", header_style)],
        [Paragraph("膜元件型号", cell_style), Paragraph(result.model_name, cell_style), Paragraph("—", cell_style)],
        [Paragraph("总膜面积", cell_style), Paragraph(_fi(result.a_actual), cell_style), Paragraph("m²", cell_style)],
        [Paragraph("总架数", cell_style), Paragraph(_fi(result.n_racks), cell_style), Paragraph("架", cell_style)],
        [Paragraph("平均通量 Javg", cell_style), Paragraph(_f(result.j_avg, 1), cell_style), Paragraph("LMH", cell_style)],
        [Paragraph("峰值通量 Jpeak", cell_style), Paragraph(_f(result.j_peak, 1), cell_style), Paragraph("LMH", cell_style)],
        [Paragraph("瞬时通量 Jinst", cell_style), Paragraph(_f(result.j_inst, 1), cell_style), Paragraph("LMH", cell_style)],
        [Paragraph("设计通量 Jdesign", cell_style), Paragraph(_f(result.j_design, 1), cell_style), Paragraph("LMH", cell_style)],
        [Paragraph("通量校核", cell_style), Paragraph(flux_status, cell_style), Paragraph("—", cell_style)],
    ]
    _add_table(story, data5)

    story.append(Spacer(1, 10))

    # 曝气系统
    story.append(Paragraph("2.2 曝气系统", h2_style))
    data6 = [
        [Paragraph("参数", header_style), Paragraph("结果", header_style), Paragraph("单位", header_style)],
        [Paragraph("总供气量", cell_style), Paragraph(_f(result.total_air_nm3min, 1), cell_style), Paragraph("Nm³/min", cell_style)],
        [Paragraph("气水比", cell_style), Paragraph(_f(result.air_water_ratio, 1), cell_style), Paragraph("—", cell_style)],
        [Paragraph("工作比", cell_style), Paragraph(_f(result.duty_cycle * 100, 1), cell_style), Paragraph("%", cell_style)],
    ]
    _add_table(story, data6)

    story.append(Spacer(1, 10))

    # 动力消耗
    story.append(Paragraph("2.3 动力消耗", h2_style))
    data7 = [
        [Paragraph("设备", header_style), Paragraph("功率", header_style), Paragraph("单位", header_style)],
        [Paragraph("曝气风机", cell_style), Paragraph(_f(result.blower_power, 1), cell_style), Paragraph("kW", cell_style)],
        [Paragraph("产水泵", cell_style), Paragraph(_f(result.pump_power, 1), cell_style), Paragraph("kW", cell_style)],
        [Paragraph("回流泵", cell_style), Paragraph(_f(result.return_pump_power, 1), cell_style), Paragraph("kW", cell_style)],
        [Paragraph("生物曝气风机", cell_style), Paragraph(_f(result.bio_blower_power, 1), cell_style), Paragraph("kW", cell_style)],
        [Paragraph("总装机功率", cell_style), Paragraph(_f(result.total_power, 1), cell_style), Paragraph("kW", cell_style)],
        [Paragraph("单位产水电耗", cell_style), Paragraph(_f(result.unit_energy, 3), cell_style), Paragraph("kWh/m³", cell_style)],
    ]
    _add_table(story, data7)

    story.append(PageBreak())

    # ---- 章节 3: 化学清洗 ----
    story.append(Paragraph("三、化学清洗药剂消耗", h1_style))
    data8 = [
        [Paragraph("参数", header_style), Paragraph("NaClO", header_style), Paragraph("柠檬酸", header_style), Paragraph("单位", header_style)],
        [Paragraph("CEB 浓度", cell_style), Paragraph(_f(inp.ceb_naclo, 0), cell_style), Paragraph(_f(inp.ceb_citric, 0), cell_style), Paragraph("mg/L", cell_style)],
        [Paragraph("CEB 频率", cell_style), Paragraph(_f(inp.ceb_freq, 1), cell_style), Paragraph(_f(inp.ceb_freq, 1), cell_style), Paragraph("次/周", cell_style)],
        [Paragraph("CIP 浓度", cell_style), Paragraph(_f(inp.cip_naclo, 0), cell_style), Paragraph(_f(inp.cip_citric, 0), cell_style), Paragraph("mg/L", cell_style)],
        [Paragraph("CIP 频率", cell_style), Paragraph(_f(inp.cip_freq, 1), cell_style), Paragraph(_f(inp.cip_freq, 1), cell_style), Paragraph("次/年", cell_style)],
        [Paragraph("年消耗量", cell_style), Paragraph(_f(result.naclo_per_year, 3), cell_style), Paragraph(_f(result.citric_per_year, 3), cell_style), Paragraph("t/a", cell_style)],
        [Paragraph("清洗水耗", cell_style), Paragraph(_f(result.wash_water_per_year, 1), cell_style), Paragraph("—", cell_style), Paragraph("m³/a", cell_style)],
    ]
    _add_table(story, data8)

    story.append(Spacer(1, 30))
    story.append(Paragraph("—— 报告结束 ——", meta_style))

    doc.build(story)
    return buf.getvalue()


def _add_table(story, data, col_widths=None):
    """简化表格添加"""
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib import colors

    if col_widths is None:
        n = len(data[0])
        # 按列数平均分配，第一列略宽
        col_widths = [180] + [130] * (n - 1) if n > 1 else [300]
    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e40af")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#f8fafc")]),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(t)


# ============================================================================
# Word 生成（python-docx）
# ============================================================================
def generate_word_report(inp: ProcessInput, result: ProcessResult,
                         project_name: str = "MBR 膜系统工艺计算书",
                         designer: str = "", design_date: str = "") -> bytes:
    """生成 Word 计算书，返回 bytes"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL

    doc = Document()

    # 设置页边距
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # 定义样式辅助函数
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)

    def add_h1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

    def add_h2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True

    def add_meta(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    def make_table(rows_data, header_fill=True):
        n_cols = len(rows_data[0])
        table = doc.add_table(rows=len(rows_data), cols=n_cols)
        table.style = "Table Grid"
        for i, row in enumerate(rows_data):
            for j, cell_text in enumerate(row):
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(9)
                        if i == 0 and header_fill:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                if i == 0 and header_fill:
                    shading = cell._tc.get_or_add_tcPr()
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "1e40af")
                    shading.append(shd)
        return table

    # 封面
    add_title(project_name)
    add_meta("MBR 膜生物反应器 - 工艺设计计算书")
    if designer:
        add_meta(f"设计人：{designer}")
    if design_date:
        add_meta(f"设计日期：{design_date}")
    add_meta(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # 设计输入
    add_h1("一、设计输入参数")
    add_h2("1.1 设计流量")
    make_table([
        ["参数", "取值", "单位"],
        ["设计流量 Q", _fi(inp.Q), "m³/d"],
        ["总变化系数 Kz", _f(inp.Kz, 2), "—"],
    ])

    add_h2("1.2 进水水质")
    make_table([
        ["项目", "数值", "单位"],
        ["COD", _f(inp.cod_in, 0), "mg/L"],
        ["BOD₅", _f(inp.bod_in, 0), "mg/L"],
        ["NH₃-N", _f(inp.nh3n_in, 0), "mg/L"],
        ["SS", _f(inp.ss_in, 0), "mg/L"],
        ["TN", _f(inp.tn_in, 0), "mg/L"],
        ["TP", _f(inp.tp_in, 0), "mg/L"],
        ["pH", _f(inp.ph_value, 1), "—"],
        ["水温", _f(inp.T, 1), "℃"],
        ["MLSS", _f(inp.MLSS, 0), "mg/L"],
    ])

    add_h2("1.3 膜组件配置")
    from mbr_calc import ALL_MODELS
    model = ALL_MODELS[inp.model_index]
    make_table([
        ["参数", "取值", "单位"],
        ["膜元件型号", model.name, "—"],
        ["单膜片面积", _f(model.sheet_area, 0), "m²"],
        ["膜孔径", _f(model.pore_size, 2), "μm"],
        ["每架膜片数", _fi(inp.sheets_per_rack), "片/架"],
        ["池数", _fi(inp.pools), "座"],
        ["每池架数", _fi(inp.racks_per_pool), "架/池"],
    ])

    add_h2("1.4 运行与清洗参数")
    make_table([
        ["参数", "取值", "单位"],
        ["基准通量 (25℃)", _f(inp.J25, 1), "LMH"],
        ["污染系数", _f(inp.fouling_factor, 2), "—"],
        ["比曝气密度 SAD", _f(inp.SAD, 0), "m³/(m²·h)"],
        ["抽吸 开/停", f"{_f(inp.suction_on, 1)} / {_f(inp.suction_off, 1)}", "min"],
        ["池内液位", _f(inp.pool_level, 1), "m"],
        ["管路损失", _f(inp.pipe_loss, 1), "m"],
        ["产水泵扬程", _f(inp.permeate_pump_head, 1), "m"],
        ["产水泵效率", _f(inp.permeate_pump_eff * 100, 0), "%"],
        ["回流比", _f(inp.return_ratio, 1), "—"],
        ["回流泵扬程", _f(inp.return_pump_head, 1), "m"],
        ["回流泵效率", _f(inp.return_pump_eff * 100, 0), "%"],
        ["风机效率", _f(inp.fan_efficiency * 100, 0), "%"],
    ])

    doc.add_page_break()

    # 计算结果
    add_h1("二、工艺计算结果")
    add_h2("2.1 膜面积与通量")
    flux_status = "✓ 在范围内" if result.j_peak <= result.j_design else "⚠ 超过设计通量"
    make_table([
        ["参数", "结果", "单位"],
        ["膜元件型号", result.model_name, "—"],
        ["总膜面积", _fi(result.a_actual), "m²"],
        ["总架数", _fi(result.n_racks), "架"],
        ["平均通量 Javg", _f(result.j_avg, 1), "LMH"],
        ["峰值通量 Jpeak", _f(result.j_peak, 1), "LMH"],
        ["瞬时通量 Jinst", _f(result.j_inst, 1), "LMH"],
        ["设计通量 Jdesign", _f(result.j_design, 1), "LMH"],
        ["通量校核", flux_status, "—"],
    ])

    add_h2("2.2 曝气系统")
    make_table([
        ["参数", "结果", "单位"],
        ["总供气量", _f(result.total_air_nm3min, 1), "Nm³/min"],
        ["气水比", _f(result.air_water_ratio, 1), "—"],
        ["工作比", _f(result.duty_cycle * 100, 1), "%"],
    ])

    add_h2("2.3 动力消耗")
    make_table([
        ["设备", "功率", "单位"],
        ["曝气风机", _f(result.blower_power, 1), "kW"],
        ["产水泵", _f(result.pump_power, 1), "kW"],
        ["回流泵", _f(result.return_pump_power, 1), "kW"],
        ["生物曝气风机", _f(result.bio_blower_power, 1), "kW"],
        ["总装机功率", _f(result.total_power, 1), "kW"],
        ["单位产水电耗", _f(result.unit_energy, 3), "kWh/m³"],
    ])

    doc.add_page_break()

    # 化学清洗
    add_h1("三、化学清洗药剂消耗")
    make_table([
        ["参数", "NaClO", "柠檬酸", "单位"],
        ["CEB 浓度", _f(inp.ceb_naclo, 0), _f(inp.ceb_citric, 0), "mg/L"],
        ["CEB 频率", _f(inp.ceb_freq, 1), _f(inp.ceb_freq, 1), "次/周"],
        ["CIP 浓度", _f(inp.cip_naclo, 0), _f(inp.cip_citric, 0), "mg/L"],
        ["CIP 频率", _f(inp.cip_freq, 1), _f(inp.cip_freq, 1), "次/年"],
        ["年消耗量", _f(result.naclo_per_year, 3), _f(result.citric_per_year, 3), "t/a"],
        ["清洗水耗", _f(result.wash_water_per_year, 1), "—", "m³/a"],
    ])

    doc.add_paragraph("—— 报告结束 ——")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
